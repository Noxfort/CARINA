# CARINA (Controlled Artificial Road-traffic Intelligence Network Architecture) is an open-source AI ecosystem for real-time, adaptive control of urban traffic light networks.
# Copyright (C) 2026 Gabriel Moraes - Noxfort Systems
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# File: src/xai/report_blocks.py
# Author: Gabriel Moraes
# Date: 2026-06-19

import os
import logging
from datetime import datetime
from typing import Dict, Any

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    logging.warning("[REPORT_BLOCKS] python-docx not installed.")

def get_translated(config: Dict[str, Any], key: str, default: str) -> str:
    lm = config.get("locale_manager")
    if lm:
        return lm.get_string(key, default=default)
    return default

class ReportBlock:
    """Base class for all modular report blocks."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        raise NotImplementedError("Subclasses must implement build().")

class HeaderBlock(ReportBlock):
    """Generates the organizational header including an optional logo and agency details."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        logo_path = config.get("logo_path")
        agency_name = config.get("agency_name", "Prefeitura Municipal / Secretaria de Trânsito")
        department_name = config.get("department_name", "Departamento de Mobilidade Inteligente")
        
        # We can add a logo to the header if it exists
        if logo_path and os.path.exists(logo_path):
            try:
                table = doc.add_table(rows=1, cols=2)
                table.autofit = True
                
                # Left cell: Logo
                cell_logo = table.cell(0, 0)
                cell_logo.width = Inches(1.5)
                p_logo = cell_logo.paragraphs[0]
                p_logo.add_run().add_picture(logo_path, width=Inches(1.2))
                
                # Right cell: Text details
                cell_text = table.cell(0, 1)
                p_text = cell_text.paragraphs[0]
                p_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run_agency = p_text.add_run(f"{agency_name.upper()}\n")
                run_agency.bold = True
                run_agency.font.size = Pt(10)
                
                run_dept = p_text.add_run(f"{department_name}\n")
                run_dept.font.size = Pt(9)
                run_dept.font.color.rgb = None
                
            except Exception as e:
                logging.error(f"[XAI_REPORT_BUILDER] Failed to add header logo table: {e}")
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_agency = p.add_run(f"{agency_name.upper()}\n")
            run_agency.bold = True
            run_agency.font.size = Pt(12)
            
            run_dept = p.add_run(department_name)
            run_dept.font.size = Pt(10)
            
        # Draw a divider line
        p_div = doc.add_paragraph()
        p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_div.add_run("―" * 60).font.color.rgb = None

class TitleBlock(ReportBlock):
    """Renders the main document title and emission timestamp."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        title_text = config.get("title", "LAUDO TÉCNICO DE EXPLICABILIDADE DE IA (XAI)")
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after = Pt(18)
        
        run_title = p_title.add_run(title_text)
        run_title.bold = True
        run_title.font.size = Pt(16)
        
        # Add Date and Time of Emission
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_date.paragraph_format.space_after = Pt(12)
        
        lbl_emission = "Data de Emissão"
        lm = config.get("locale_manager")
        if lm:
            lbl_emission = lm.get_string("xai_viewer.emission_date", default="Data de Emissão")
            if "emission_date" not in lm.current_lang_data.get("xai_viewer", {}):
                # Fallback to mfd_viewer if needed
                lbl_emission = lm.get_string("mfd_viewer.emission_date", default=lbl_emission)
                
        run_date = p_date.add_run(f"{lbl_emission}: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        run_date.font.size = Pt(9.5)
        run_date.italic = True

class MetadataBlock(ReportBlock):
    """Creates a structured metadata table with information regarding the agent or run under analysis."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        agent_id = context.get("agent_id", "UNKNOWN")
        scenario = context.get("scenario", "Live Session")
        engine_ver = context.get("engine_version", "CARINA v1.0.0")
        
        mode = config.get("mode", "XAI")
        
        if mode == "MFD":
            heading = get_translated(config, "structured_report.mfd_metadata_title", "1. Detalhes da Análise de Desempenho MFD")
            doc.add_heading(heading, level=1)
            
            headers = [
                (get_translated(config, "structured_report.mfd_label_scenario", "Cenário de Simulação:"), scenario),
                (get_translated(config, "structured_report.mfd_label_engine", "Motor Analítico MFD:"), engine_ver),
                (get_translated(config, "structured_report.mfd_label_control", "Sistema de Controle:"), agent_id)
            ]
        else:
            heading = get_translated(config, "structured_report.xai_metadata_title", "1. Detalhes Gerais da Sessão")
            doc.add_heading(heading, level=1)
            
            headers = [
                (get_translated(config, "structured_report.xai_label_agent", "Identificador do Agente:"), agent_id),
                (get_translated(config, "structured_report.xai_label_scenario", "Cenário de Operação:"), scenario),
                (get_translated(config, "structured_report.xai_label_engine", "Motor Analítico:"), engine_ver)
            ]
            
        table = doc.add_table(rows=len(headers), cols=2)
        table.style = 'Table Grid'
        
        for idx, (label, val) in enumerate(headers):
            cell_label = table.cell(idx, 0)
            cell_label.paragraphs[0].add_run(label).bold = True
            cell_label.paragraphs[0].runs[0].font.size = Pt(10)
            
            cell_val = table.cell(idx, 1)
            cell_val.paragraphs[0].add_run(val)
            cell_val.paragraphs[0].runs[0].font.size = Pt(10)
            
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

class ChartBlock(ReportBlock):
    """Integrates the matplotlib plot representing importance attributions or performance."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        image_path = context.get("image_path")
        
        if image_path and os.path.exists(image_path):
            mode = config.get("mode", "XAI")
            if mode == "MFD":
                heading = get_translated(config, "structured_report.mfd_chart_title", "2. Visualização do Diagrama Fundamental Macroscópico (MFD)")
                caption = get_translated(config, "structured_report.mfd_chart_caption", "Figura 1: Curva de otimização MFD exibindo produção versus acumulação da malha viária.")
            else:
                heading = get_translated(config, "structured_report.xai_chart_title", "2. Análise de Atribuição Matemática (Captum)")
                caption = get_translated(config, "structured_report.xai_chart_caption", "Figura 1: Importância relativa dos sensores de aproximação calculada via gradientes integrados.")
                
            doc.add_heading(heading, level=1)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p_img.add_run().add_picture(image_path, width=Inches(5.5))
                
                # Caption
                p_caption = doc.add_paragraph()
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_caption.paragraph_format.space_after = Pt(12)
                run_caption = p_caption.add_run(caption)
                run_caption.font.size = Pt(8.5)
                run_caption.italic = True
            except Exception as e:
                logging.error(f"[XAI_REPORT_BUILDER] Failed to insert chart: {e}")
                err_text = get_translated(config, "structured_report.error_loading_chart", "[ERRO: Não foi possível carregar a imagem do gráfico.]")
                doc.add_paragraph(err_text)
        else:
            logging.info("[XAI_REPORT_BUILDER] Chart image path not found or empty.")

class ContentBlock(ReportBlock):
    """Processes the natural language explanation returned by the LLM, translating markdown headings."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        text_content = context.get("text_content", "")
        mode = config.get("mode", "XAI")
        
        if mode == "MFD":
            heading = get_translated(config, "structured_report.mfd_content_title", "3. Relatório de Desempenho e Avaliação de Otimização")
            fallback_text = get_translated(config, "structured_report.mfd_content_fallback", "Nenhum descritivo semântico foi gerado pelo modelo de otimização.")
        else:
            heading = get_translated(config, "structured_report.xai_content_title", "3. Descritivo Semântico e Parecer Técnico")
            fallback_text = get_translated(config, "structured_report.xai_content_fallback", "Nenhum descritivo semântico foi gerado pelo modelo de linguagem.")
            
        doc.add_heading(heading, level=1)
        
        if not text_content:
            doc.add_paragraph(fallback_text)
            return

        for line in text_content.split('\n'):
            line_str = line.strip()
            if not line_str:
                continue
                
            if line_str.startswith('###'):
                p = doc.add_paragraph()
                run = p.add_run(line_str.strip('#* '))
                run.bold = True
                run.font.size = Pt(11.5)
                p.paragraph_format.space_before = Pt(6)
            elif line_str.startswith('##') or line_str.startswith('**') and line_str.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line_str.strip('#* '))
                run.bold = True
                run.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(8)
            else:
                doc.add_paragraph(line_str)

class SignatureBlock(ReportBlock):
    """Draws a professional signature block at the bottom for official validation."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        secretary_name = config.get("secretary_name", "Dr. Gabriel Moraes")
        secretary_title = config.get("secretary_title", "Secretário de Mobilidade e Trânsito")
        mode = config.get("mode", "XAI")
        
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(36)
        
        p_sig = doc.add_paragraph()
        p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_line = p_sig.add_run("___________________________________________________\n")
        run_line.bold = True
        
        run_name = p_sig.add_run(f"{secretary_name}\n")
        run_name.bold = True
        run_name.font.size = Pt(11)
        
        run_title = p_sig.add_run(secretary_title)
        run_title.font.size = Pt(9.5)
        
        # Conformity footer
        p_conf = doc.add_paragraph()
        p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_conf.paragraph_format.space_before = Pt(24)
        
        if mode == "MFD":
            conf_text = get_translated(
                config, 
                "structured_report.mfd_conformity_text", 
                "Este laudo foi gerado de forma determinística pelo motor de otimização CARINA MFD. Ele atesta as métricas de performance macroscópica da rede de tráfego analisada."
            )
        else:
            conf_text = get_translated(
                config, 
                "structured_report.xai_conformity_text", 
                "Este documento foi consolidado de forma determinística pelo motor Neuro-Simbólico CARINA XAI. Ele atesta as condições puras da leitura de topologia e do comportamento da rede neural."
            )
            
        run_conf = p_conf.add_run(conf_text)
        run_conf.font.size = Pt(8)
        run_conf.font.color.rgb = None
