import os
import logging
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    logging.warning("[REPORT_GEN] python-docx não instalado. Relatórios não poderão ser gerados em .docx")

from xai.semantic_transducer import SemanticTransducer
from utils.settings_manager import SettingsManager

class ReportGenerator:
    """
    Gera laudos técnicos profissionais em .docx usando o SemanticTransducer (Qwen LLM).
    O processo rodará com Temperature=0.0 de forma nativa para garantir a integridade.
    """
    def __init__(self, locale_manager=None):
        self.locale_manager = locale_manager
        
        # O Transducer precisa saber o caminho exato do GGUF.
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.model_path = os.path.join(base_dir, "Model_Vault", "Qwen3.5-2B-UD-Q6_K_XL.gguf")
        
    def generate_docx_report(self, analysis_results: dict, scenario_dir: str, net_file_path: str):
        if not analysis_results:
            logging.warning("[REPORT_GEN] Sem dados para gerar relatório.")
            return None
            
        logging.info("[REPORT_GEN] Iniciando geração do Laudo Técnico Profissional via LLM...")
        
        try:
            # 1. Instancia o LLM garantindo isolamento
            transducer = SemanticTransducer(self.model_path, use_gpu=False, offload_to_cpu=True)
            transducer.load_resources()
            
            try:
                settings_manager = SettingsManager()
                ui_language = settings_manager.get_setting('General', 'language', 'pt_br')
            except Exception:
                ui_language = 'pt_br'

            # 2. Prepara os dados brutos para injetar no Qwen
            input_data = {
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "mode": "LAUDO_ESTATISTICO",
                "language": ui_language,
                "attributions": analysis_results
            }
            
            # 3. Executa a inferência. O Transducer nativamente já força temperature=0.0!
            logging.info("[REPORT_GEN] Transducer extraindo contexto semântico (Temperature=0.0)...")
            report_text = transducer.generate_report(input_data)
            
            # 4. Formata o Documento Word (.docx)
            doc = Document()
            
            # Título principal
            title = doc.add_heading('LAUDO DESCRITIVO DE TRÁFEGO E INFRAESTRUTURA', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtítulo com data
            p_date = doc.add_paragraph()
            p_date.add_run(f"Data de Emissão: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}").bold = True
            p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Adiciona o mapa se existir
            map_path = os.path.join(scenario_dir, "map_planning.png")
            if os.path.exists(map_path):
                doc.add_heading('1. Mapa de Planejamento Tático', level=1)
                doc.add_picture(map_path, width=Inches(6.0))
                
            # Adiciona a resposta do LLM formatada
            doc.add_heading('2. Descritivo Operacional (Organização de Dados)', level=1)
            
            # Parser básico de markdown do LLM para o Word
            for line in report_text.split('\n'):
                if line.strip().startswith('##') or line.strip().startswith('**'):
                    p = doc.add_paragraph()
                    p.add_run(line.strip('#* ')).bold = True
                elif line.strip():
                    doc.add_paragraph(line.strip())
            
            # Rodapé de conformidade
            doc.add_page_break()
            footer = doc.add_paragraph()
            r = footer.add_run("Este documento foi consolidado de forma determinística e imparcial (Temperature=0.0) pelo motor Neuro-Simbólico CARINA XAI. Ele atesta as condições puras da leitura de topologia.")
            r.font.size = Pt(8)
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Salva o arquivo final
            docx_path = os.path.join(scenario_dir, "Laudo_Tecnico_Oficial.docx")
            doc.save(docx_path)
            
            logging.info(f"[REPORT_GEN] Laudo em DOCX gerado com sucesso em: {docx_path}")
            return docx_path
            
        except Exception as e:
            logging.error(f"[REPORT_GEN] Erro ao gerar o arquivo .docx: {e}", exc_info=True)
            return None
