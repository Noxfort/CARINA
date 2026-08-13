# CARINA (Controlled Artificial Road-traffic Intelligence Network Architecture) is an open-source AI ecosystem for real-time, adaptive control of urban traffic light networks.
# Copyright (C) 2026 Gabriel Moraes - Noxfort Systems
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# File: src/blocks/content.py
# Author: Gabriel Moraes
# Date: 2026-07-02

from typing import Dict, Any

from .base import ReportBlock, get_translated, add_markdown_paragraph

try:
    from docx.shared import RGBColor
except ImportError:
    pass

class ContentBlock(ReportBlock):
    """Processes the natural language explanation returned by the LLM, translating markdown headings."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        text_content = context.get("text_content", "")
        mode = config.get("mode", "XAI")
        
        content_title = config.get("content_title")
        content_fallback = config.get("content_fallback")

        should_add_heading = True
        if text_content and ("## " in text_content or "# " in text_content or "### " in text_content):
            should_add_heading = False
            fallback_text = "Nenhum laudo analítico disponível."
        elif content_title is not None and content_fallback is not None:
            heading = content_title
            fallback_text = content_fallback
        elif mode == "PLANNING":
            should_add_heading = False
            fallback_text = "Nenhum laudo analítico de planejamento disponível."
        elif mode == "MFD":
            heading = get_translated(config, "structured_report.mfd_content_title", "Relatório de Desempenho e Avaliação de Otimização")
            fallback_text = get_translated(config, "structured_report.mfd_content_fallback", "Nenhum descritivo semântico foi gerado pelo modelo de otimização.")
        else:
            heading = get_translated(config, "structured_report.xai_content_title", "Descritivo Semântico e Parecer Técnico")
            fallback_text = get_translated(config, "structured_report.xai_content_fallback", "Nenhum descritivo semântico foi gerado pelo modelo de linguagem.")
            
        if should_add_heading:
            h = doc.add_heading(heading, level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        if not text_content:
            p = doc.add_paragraph()
            run = p.add_run(fallback_text)
            run.font.color.rgb = RGBColor(0, 0, 0)
            return

        font_name = config.get("font_name", "Arial")
        font_size = float(config.get("font_size", 11.0))
        
        from blocks.markdown_to_docx import render_markdown_to_docx
        render_markdown_to_docx(doc, text_content, default_font_size=font_size, font_name=font_name, context=context, config=config)
