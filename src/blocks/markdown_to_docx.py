# CARINA (Controlled Artificial Road-traffic Intelligence Network Architecture) is an open-source AI ecosystem for real-time, adaptive control of urban traffic light networks.
# Copyright (C) 2026 Gabriel Moraes - Noxfort Systems
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# File: src/blocks/markdown_to_docx.py
# Author: Gabriel Moraes
# Date: July 25, 2026

import re
from typing import Dict, Any

try:
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

from blocks.math_cleaner import clean_latex_math
from blocks.docx_text_builder import add_formatted_text_to_paragraph, add_omml_equation_to_document
from blocks.docx_table_builder import render_markdown_table

def render_markdown_to_docx(doc: Any, markdown_text: str, default_font_size: float = 11.0, font_name: str = "Arial", context: Dict[str, Any] = None, config: Dict[str, Any] = None) -> None:
    """
    Renders standard Markdown text into a python-docx Document object.
    Orchestrates:
    - Headers (# H1, ## H2, ### H3, #### H4)
    - Code blocks (```)
    - Markdown tables (| Col1 | Col2 |) via docx_table_builder
    - Centralized equations ($$...$$) and ABNT NBR 14724 tag alignment
    - Selective bullet and numbered lists
    - Automatic page breaks before ANEXO I and signature placement
    """
    if context is None:
        context = {}
    if config is None:
        config = {}

    if not markdown_text:
        return

    lines = markdown_text.split('\n')
    i = 0
    in_code_block = False

    NON_BULLET_LABELS = [
        "cenário de simulação:", "motor analítico mfd:", "sistema de controle:",
        "identificador do agente:", "cenário de operação:", "motor analítico:",
        "atraso médio (p95):", "extensão da fila:", "taxa de saturação:",
        "o que são as variáveis", "explicação para gestão pública", "o que é o conceito matemático"
    ]

    last_was_page_break = False

    while i < len(lines):
        line = lines[i]
        line_str = line.strip()

        # 0. Handle Code Block Fences (```)
        if line_str.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if not line_str:
            i += 1
            continue

        # 0.1 Handle explicit Page Break tags (<pagebreak>, <!-- PAGE BREAK -->, [page_break])
        if "<pagebreak>" in line_str.lower() or "<!-- page break -->" in line_str.lower() or "[page_break]" in line_str.lower():
            doc.add_page_break()
            last_was_page_break = True
            cleaned_line = re.sub(r"(?i)<pagebreak>|<!-- page break -->|\[page_break\]", "", line_str).strip()
            if not cleaned_line:
                i += 1
                continue
            line_str = cleaned_line

        # 0.2 Handle explicit Signature Block tags ([signature_block], [signature])
        if "[signature_block]" in line_str.lower() or "[signature]" in line_str.lower():
            try:
                from blocks.signature import SignatureBlock
                SignatureBlock().build(doc, context, config)
            except Exception:
                pass
            cleaned_line = re.sub(r"(?i)\[signature_block\]|\[signature\]", "", line_str).strip()
            if not cleaned_line:
                i += 1
                continue
            line_str = cleaned_line

        # If inside a code block, format line cleanly as an indented block paragraph
        if in_code_block:
            last_was_page_break = False
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text_to_paragraph(p, line_str, font_size=default_font_size - 0.5, font_name="Courier New")
            i += 1
            continue

        # 1. Handle Markdown Tables (| Col1 | Col2 |)
        if line_str.startswith('|'):
            last_was_page_break = False
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            render_markdown_table(doc, table_lines, font_name=font_name, font_size=default_font_size)
            continue

        # 2. Headers (# H1, ## H2, ### H3, #### H4)
        if line_str.startswith('#'):
            level = 0
            for char in line_str:
                if char == '#':
                    level += 1
                else:
                    break

            header_text = line_str[level:].strip()
            header_text = clean_latex_math(header_text)

            # Render Signature Block before Anexo or Apêndice sections if not already rendered
            if "ANEXO" in header_text.upper() or "APÊNDICE" in header_text.upper():
                if not context.get("_signature_rendered", False):
                    try:
                        from blocks.signature import SignatureBlock
                        SignatureBlock().build(doc, context, config)
                    except Exception:
                        pass
                if not last_was_page_break:
                    doc.add_page_break()
                    last_was_page_break = True
            else:
                last_was_page_break = False

            # Map level to professional font sizes
            if level == 1:
                font_size = 14.0
                space_before = 12
            elif level == 2:
                font_size = 12.5
                space_before = 10
            elif level == 3:
                font_size = 11.5
                space_before = 8
            else:
                font_size = 11.0
                space_before = 6

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True

            add_formatted_text_to_paragraph(p, header_text, font_size=font_size, bold=True, font_name=font_name)
            i += 1
            continue

        # 3. Centralized Equations (lines wrapped in $$) - ABNT NBR 14724 Standard
        if line_str.startswith("$$") and line_str.endswith("$$"):
            eq_raw = line_str[2:-2].strip()

            tag_match = re.search(r'\\tag\{(\d+)\}', eq_raw) or re.search(r'\s*\((\d+)\)\s*$', eq_raw)
            tag_str = ""
            if tag_match:
                tag_num = tag_match.group(1)
                tag_str = f"({tag_num})"
                eq_raw = re.sub(r'\\tag\{(\d+)\}', '', eq_raw)
                eq_raw = re.sub(r'\s*\(\d+\)\s*$', '', eq_raw).strip()

            # Try native Word OMML fraction rendering first
            if not add_omml_equation_to_document(doc, eq_raw, tag_str, font_name=font_name, font_size=default_font_size):
                cleaned_eq = clean_latex_math(eq_raw)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)

                if tag_str:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    add_formatted_text_to_paragraph(p, f"      **{cleaned_eq}**", font_size=default_font_size, font_name=font_name)
                    r_tag = p.add_run(f"\t\t{tag_str}")
                    r_tag.bold = True
                    r_tag.font.name = font_name
                    r_tag.font.size = Pt(default_font_size)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_formatted_text_to_paragraph(p, f"**{cleaned_eq}**", font_size=default_font_size, font_name=font_name)

            i += 1
            continue

        # 4. Bullet list and selective filtering
        if line_str.startswith('- ') or line_str.startswith('* '):
            list_text = line_str[2:].strip()
            clean_check = list_text.replace("**", "").strip().lower()

            is_non_bullet = any(clean_check.startswith(lbl) for lbl in NON_BULLET_LABELS)
            if not is_non_bullet and len(list_text) > 130 and ("justificativ" in clean_check or "justification" in clean_check or "обоснование" in clean_check or "论证" in clean_check):
                is_non_bullet = True

            if is_non_bullet or clean_check.startswith("o que são") or clean_check.startswith("explicação para") or clean_check.startswith("o que é"):
                line_str = list_text
            else:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                list_text = clean_latex_math(list_text)
                add_formatted_text_to_paragraph(p, list_text, font_size=default_font_size, font_name=font_name)
                i += 1
                continue

        # 5. Numbered list
        match_num = re.match(r'^(\d+)\.\s(.*)', line_str)
        if match_num:
            list_text = match_num.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            list_text = clean_latex_math(list_text)
            add_formatted_text_to_paragraph(p, list_text, font_size=default_font_size, font_name=font_name)
            i += 1
            continue

        # 6. Standard Paragraph (Official Redaction Formatting)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5

        if len(line_str) > 80 and not line_str.startswith("**"):
            p.paragraph_format.first_line_indent = Cm(1.25)

        cleaned_line = clean_latex_math(line_str)
        add_formatted_text_to_paragraph(p, cleaned_line, font_size=default_font_size, font_name=font_name)
        i += 1
