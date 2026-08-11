# CARINA (Controlled Artificial Road-traffic Intelligence Network Architecture) is an open-source AI ecosystem for real-time, adaptive control of urban traffic light networks.
# Copyright (C) 2026 Gabriel Moraes - Noxfort Systems
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# File: src/blocks/title.py
# Author: Gabriel Moraes
# Date: 2026-07-02

from datetime import datetime
from typing import Dict, Any

from .base import ReportBlock

try:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

class TitleBlock(ReportBlock):
    """Renders the main document title and emission timestamp."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        title_text = str(config.get("title") or "LAUDO TÉCNICO DE ENGENHARIA DE TRÂNSITO")
        
        # Official redaction date mask (e.g., Apucarana - PR, 24 de julho de 2026)
        months_pt = [
            "janeiro", "fevereiro", "março", "abril", "maio", "junho",
            "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
        ]
        now = datetime.now()
        month_str = months_pt[now.month - 1]
        
        city = config.get("city") or config.get("report_city") or "Apucarana"
        state_uf = config.get("state_uf") or config.get("report_state_uf") or "PR"
        protocol_num = config.get("protocol_number") or config.get("report_protocol_number") or "042/2026"
        
        official_date_str = f"{city} - {state_uf}, {now.day} de {month_str} de {now.year}"

        # 1. Date (Right aligned)
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_date.paragraph_format.space_after = Pt(12)
        run_date = p_date.add_run(official_date_str)
        run_date.font.size = Pt(10)
        run_date.italic = True
        run_date.font.color.rgb = RGBColor(0, 0, 0)

        # 2. Main Protocol Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after = Pt(6)
        
        clean_title = title_text.replace("|", "").strip().upper()
        if "Nº" not in clean_title and "NO" not in clean_title:
            clean_title = f"{clean_title} Nº {protocol_num}"
            
        run_title = p_title.add_run(clean_title)
        run_title.bold = True
        run_title.font.size = Pt(14)
        run_title.font.color.rgb = RGBColor(0, 0, 0)

        # 3. Official Ementa (Right indented block)
        p_ementa = doc.add_paragraph()
        p_ementa.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_ementa.paragraph_format.left_indent = Pt(180) # Indent to the right half
        p_ementa.paragraph_format.space_before = Pt(6)
        p_ementa.paragraph_format.space_after = Pt(18)
        
        ementa_text = config.get("ementa_text") or "Assunto: Análise da Capacidade Operacional, Avaliação de Warrants Técnicos (CONTRAN/MUTCD) e Recomendação Semafórica para a Malha Viária Urbana."
        run_ementa = p_ementa.add_run(ementa_text)
        run_ementa.font.size = Pt(9.5)
        run_ementa.italic = True
        run_ementa.font.color.rgb = RGBColor(60, 60, 60)
