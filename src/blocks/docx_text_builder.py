# CARINA (Controlled Artificial Road-traffic Intelligence Network Architecture) is an open-source AI ecosystem for real-time, adaptive control of urban traffic light networks.
# Copyright (C) 2026 Gabriel Moraes - Noxfort Systems
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# File: src/blocks/docx_text_builder.py
# Author: Gabriel Moraes
# Date: August 9, 2026

import os
import re
import json
import logging
from typing import Dict, Any, List

try:
    from docx.shared import Pt, RGBColor
    from docx.oxml import parse_xml
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

_cached_omml_config: Dict[str, Any] = None
_cached_subscript_config: Dict[str, Any] = None

def _load_omml_templates() -> Dict[str, Any]:
    """Dynamically loads OMML equation XML templates from config/omml_equation_templates.json with in-memory caching."""
    global _cached_omml_config
    if _cached_omml_config is not None:
        return _cached_omml_config

    base_dir = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        os.path.join(base_dir, "..", "..", "config", "omml_equation_templates.json"),
        os.path.join(base_dir, "..", "config", "omml_equation_templates.json"),
        os.path.join(os.getcwd(), "config", "omml_equation_templates.json")
    ]

    for json_path in candidates:
        if os.path.exists(json_path):
            try:
                with open(json_path, "r", encoding="utf-8") as f:
                    _cached_omml_config = json.load(f)
                    return _cached_omml_config
            except Exception as e:
                logging.warning(f"Failed to load OMML templates from '{json_path}': {e}")

    _cached_omml_config = {
        "namespaces": 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
        "templates": {}
    }
    return _cached_omml_config

def _load_subscript_rules() -> Dict[str, Any]:
    """Dynamically loads subscript regex rules from config/subscript_rules.json with in-memory caching."""
    global _cached_subscript_config
    if _cached_subscript_config is not None:
        return _cached_subscript_config

    base_dir = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        os.path.join(base_dir, "..", "..", "config", "subscript_rules.json"),
        os.path.join(base_dir, "..", "config", "subscript_rules.json"),
        os.path.join(os.getcwd(), "config", "subscript_rules.json")
    ]

    for json_path in candidates:
        if os.path.exists(json_path):
            try:
                with open(json_path, "r", encoding="utf-8") as f:
                    _cached_subscript_config = json.load(f)
                    return _cached_subscript_config
            except Exception as e:
                logging.warning(f"Failed to load subscript rules from '{json_path}': {e}")

    _cached_subscript_config = {
        "subscript_regex_pattern": r'\b([a-zA-Z]{1,3})_\{?([a-zA-Z0-9]+)\}?|\b(v)(real|limite)\b|\b(F)(ideal)\b|\b(P)(95)\b'
    }
    return _cached_subscript_config

def add_omml_equation_to_document(doc: Any, eq_raw: str, tag_str: str = "", font_name: str = "Arial", font_size: float = 11.0) -> bool:
    """
    Renders native Word OMML equations with stacked vertical fractions (horizontal bar)
    and right-aligned ABNT tag (1)-(4), dynamically driven by JSON configuration.
    """
    try:
        from docx.oxml import parse_xml
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return False

    omml_config = _load_omml_templates()
    ns_decl = omml_config.get("namespaces", 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')
    templates_dict = omml_config.get("templates", {})

    eq_clean = eq_raw.replace("$", "").strip()
    omml_xml = None

    # Check JSON configured equation templates by keywords
    for tpl_name, tpl_info in templates_dict.items():
        keywords = tpl_info.get("keywords", [])
        if any(kw in eq_clean for kw in keywords):
            raw_xml_str = tpl_info.get("xml", "")
            omml_xml = raw_xml_str.replace("{ns_decl}", ns_decl)
            break

    # Generic Division / Fraction fallback (num / den)
    if not omml_xml and ('/' in eq_clean or r'\frac' in eq_clean):
        if '=' in eq_clean:
            left_part, right_part = eq_clean.split('=', 1)
            left_part = left_part.strip() + " = "
        else:
            left_part = ""
            right_part = eq_clean

        if '/' in right_part:
            num_str, den_str = right_part.split('/', 1)
        else:
            num_str, den_str = right_part, ""
        
        num_clean = num_str.strip()
        den_clean = den_str.strip().strip('()')
        if not den_clean:
            den_clean = "1"

        omml_xml = (
            f'<m:oMathPara {ns_decl}>\n'
            f'  <m:oMath>\n'
            f'    <m:r><m:t>{left_part}</m:t></m:r>\n'
            f'    <m:f>\n'
            f'      <m:num><m:r><m:t>{num_clean}</m:t></m:r></m:num>\n'
            f'      <m:den><m:r><m:t>{den_clean}</m:t></m:r></m:den>\n'
            f'    </m:f>\n'
            f'  </m:oMath>\n'
            f'</m:oMathPara>'
        )

    if not omml_xml:
        return False

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if tag_str else WD_ALIGN_PARAGRAPH.CENTER

    omml_elem = parse_xml(omml_xml)
    p._element.append(omml_elem)

    if tag_str:
        r_tag = p.add_run(f"\t\t{tag_str}")
        r_tag.bold = True
        r_tag.font.name = font_name
        r_tag.font.size = Pt(font_size)

    return True

def add_formatted_text_to_paragraph(p: Any, text: str, font_size: float = 11.0, bold: bool = False, font_name: str = "Arial") -> None:
    """Splits text by ** (bold) and * (italic) to apply appropriate formatting runs inline."""
    parts = text.split('**')
    is_bold_part = False
    for part in parts:
        if part:
            sub_parts = part.split('*')
            is_italic_sub = False
            for sub in sub_parts:
                if sub:
                    add_text_run_with_subscript_support(p, sub, font_size=font_size, is_bold=(is_bold_part or bold), is_italic=is_italic_sub, font_name=font_name)
                is_italic_sub = not is_italic_sub
        is_bold_part = not is_bold_part

def add_text_run_with_subscript_support(p: Any, text_segment: str, font_size: float = 11.0, is_bold: bool = False, is_italic: bool = False, font_name: str = "Arial") -> None:
    """
    Adds runs to paragraph p, detecting variable subscripts driven by JSON subscript configuration
    and applying native Word subscript (run.font.subscript = True) to the subscript portion.
    """
    sub_config = _load_subscript_rules()
    pattern = sub_config.get(
        "subscript_regex_pattern",
        r'\b([a-zA-Z]{1,3})_\{?([a-zA-Z0-9]+)\}?|\b(v)(real|limite)\b|\b(F)(ideal)\b|\b(P)(95)\b'
    )

    last_idx = 0
    for match in re.finditer(pattern, text_segment):
        start, end = match.span()
        # Add preceding normal text
        if start > last_idx:
            normal_part = text_segment[last_idx:start]
            r = p.add_run(normal_part)
            r.font.name = font_name
            r.font.size = Pt(font_size)
            try:
                r.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass
            if is_bold:
                r.bold = True
            if is_italic:
                r.italic = True

        g1, g2 = match.group(1), match.group(2)
        if g1 and g2:
            base_var = g1
            sub_text = g2
        else:
            # Captures for (v)(real|limite), (F)(ideal), (P)(95)
            groups = [g for g in match.groups() if g is not None]
            if len(groups) >= 2:
                base_var = groups[-2]
                sub_text = groups[-1]
            else:
                base_var = text_segment[start:end]
                sub_text = ""

        # Add base variable (e.g. 'v', 'F', 'P')
        r_base = p.add_run(base_var)
        r_base.font.name = font_name
        r_base.font.size = Pt(font_size)
        try:
            r_base.font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            pass
        if is_bold:
            r_base.bold = True
        if is_italic:
            r_base.italic = True

        # Add subscript part (e.g. 'real', 'limite', 'ideal', '95')
        r_sub = p.add_run(sub_text)
        r_sub.font.name = font_name
        r_sub.font.size = Pt(font_size)
        try:
            r_sub.font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            pass
        if is_bold:
            r_sub.bold = True
        if is_italic:
            r_sub.italic = True
        r_sub.font.subscript = True

        last_idx = end

    # Add remaining text
    if last_idx < len(text_segment):
        rem_part = text_segment[last_idx:]
        r = p.add_run(rem_part)
        r.font.name = font_name
        r.font.size = Pt(font_size)
        try:
            r.font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            pass
        if is_bold:
            r.bold = True
        if is_italic:
            r.italic = True
