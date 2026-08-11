# CARINA (Controlled Artificial Road-traffic Intelligence Network Architecture) is an open-source AI ecosystem for real-time, adaptive control of urban traffic light networks.
# Copyright (C) 2026 Gabriel Moraes - Noxfort Systems
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# File: src/blocks/signature.py
# Author: Gabriel Moraes
# Date: 2026-07-02

from typing import Dict, Any

from .base import ReportBlock, get_translated

try:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

class SignatureBlock(ReportBlock):
    """Draws a professional signature block at the bottom for official validation."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        if context.get("_signature_rendered", False):
            return
        context["_signature_rendered"] = True

        secretary_name = config.get("secretary_name", "Dr. Gabriel Moraes")
        secretary_title = config.get("secretary_title", "Secretário de Mobilidade e Trânsito")
        agency_name = config.get("agency_name") or config.get("report_agency_name") or "Prefeitura Municipal de Apucarana"
        
        ord_enabled = str(config.get("ordinance_enabled") or config.get("report_ordinance_enabled", "")).lower() in ["true", "1", "yes"]
        ord_number = str(config.get("ordinance_number") or config.get("report_ordinance_number", "")).strip()
        
        mode = config.get("mode", "XAI")
        
        clean_name = str(secretary_name).replace("|", "").strip()
        clean_title = str(secretary_title).replace("|", "").strip()
        clean_agency = str(agency_name).replace("|", "").strip()

        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(36)
        
        # 1. Signature Rule Line
        p_line = doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(4)
        run_line = p_line.add_run("___________________________________________________")
        run_line.bold = True
        run_line.font.color.rgb = RGBColor(0, 0, 0)
        
        # 2. Signatory Name
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_before = Pt(0)
        p_name.paragraph_format.space_after = Pt(2)
        run_name = p_name.add_run(clean_name)
        run_name.bold = True
        run_name.font.size = Pt(11)
        run_name.font.color.rgb = RGBColor(0, 0, 0)
        
        # 3. Signatory Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(2)
        run_title = p_title.add_run(clean_title)
        run_title.font.size = Pt(9.5)
        run_title.font.color.rgb = RGBColor(0, 0, 0)
        
        # 4. Ordinance Line (if enabled)
        if ord_enabled and ord_number:
            p_ord = doc.add_paragraph()
            p_ord.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_ord.paragraph_format.space_before = Pt(0)
            p_ord.paragraph_format.space_after = Pt(12)
            
            if not ord_number.lower().startswith("portaria"):
                ord_line_str = f"Portaria nº {ord_number} – {clean_agency}"
            else:
                ord_line_str = f"{ord_number} – {clean_agency}"
                
            run_ord = p_ord.add_run(ord_line_str)
            run_ord.font.size = Pt(9.5)
            run_ord.font.color.rgb = RGBColor(0, 0, 0)
        
        p_conf = doc.add_paragraph()
        p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_conf.paragraph_format.space_before = Pt(24)
        
        conformity_text = config.get("conformity_text")
        if conformity_text is not None:
            conf_text = conformity_text
        elif mode == "MFD":
            conf_text = get_translated(
                config, 
                "structured_report.mfd_conformity_text", 
                "Este laudo foi gerado de forma determinística pelo motor de otimização CARINA MFD. Ele atesta as métricas de performance macroscópica da rede de tráfego analisada."
            )
        else:
            conf_text = get_translated(
                config, 
                "structured_report.xai_conformity_text", 
                "Este documento foi consolidado de forma determinística pelo motor Neuro-Simbólico CARINA XAI. Ele atesta as condições puras da leitura de topologia e do comportamento da rede neural."
            )
            
        run_conf = p_conf.add_run(conf_text)
        run_conf.font.size = Pt(8)
        run_conf.font.color.rgb = RGBColor(0, 0, 0)
        context["_signature_rendered"] = True
