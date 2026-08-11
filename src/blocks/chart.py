# CARINA (Controlled Artificial Road-traffic Intelligence Network Architecture) is an open-source AI ecosystem for real-time, adaptive control of urban traffic light networks.
# Copyright (C) 2026 Gabriel Moraes - Noxfort Systems
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# File: src/blocks/chart.py
# Author: Gabriel Moraes
# Date: 2026-07-02

import os
import logging
from typing import Dict, Any

from .base import ReportBlock, get_translated

try:
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

class ChartBlock(ReportBlock):
    """Integrates the matplotlib plot representing importance attributions or performance."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        image_path = context.get("image_path")
        mode = config.get("mode", "XAI")
        
        # Fallback: if image_path is not specified, does not exist, or is empty, try locating it via results_dir
        if not image_path or not os.path.exists(image_path) or os.path.getsize(image_path) == 0:
            results_dir = context.get("results_dir")
            if results_dir:
                if mode == "PLANNING":
                    fallback_paths = [
                        os.path.join(results_dir, "map_planning.png"),
                        os.path.join(results_dir, "maps", "map_planning.png")
                    ]
                elif mode == "MFD":
                    fallback_paths = [
                        os.path.join(results_dir, "mfd_curve.png"),
                        os.path.join(results_dir, "plots", "mfd_curve.png")
                    ]
                else:
                    fallback_paths = [
                        os.path.join(results_dir, "xai_importance.png"),
                        os.path.join(results_dir, "plots", "xai_importance.png")
                    ]
                for p in fallback_paths:
                    if os.path.exists(p) and os.path.getsize(p) > 0:
                        image_path = p
                        break
        
        if image_path and os.path.exists(image_path) and os.path.getsize(image_path) > 0:
            chart_title = config.get("chart_title")
            chart_caption = config.get("chart_caption")
            if chart_title is not None and chart_caption is not None:
                heading = chart_title
                caption = chart_caption
            elif mode == "MFD":
                heading = get_translated(config, "structured_report.mfd_chart_title", "2. Visualização do Diagrama Fundamental Macroscópico (MFD)")
                caption = get_translated(config, "structured_report.mfd_chart_caption", "Figura 1: Curva de otimização MFD exibindo produção versus acumulação da malha viária.")
            else:
                heading = get_translated(config, "structured_report.xai_chart_title", "2. MAPA DE PLANEJAMENTO TÁTICO DA MALHA VIÁRIA")
                caption = get_translated(config, "structured_report.xai_chart_caption", "Figura 1 – Mapa com Recomendações Espaciais da Malha Viária.")
                
            # Section Heading (Level 1)
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            run_h = h.add_run(heading.replace("|", "").strip())
            run_h.bold = True
            run_h.font.size = Pt(12)
            run_h.font.color.rgb = RGBColor(0, 0, 0)
            
            # ABNT Figure Title (Above Image)
            p_fig_title = doc.add_paragraph()
            p_fig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_fig_title.paragraph_format.space_before = Pt(4)
            p_fig_title.paragraph_format.space_after = Pt(4)
            
            clean_caption = caption.replace("|", "").strip()
            if not clean_caption.lower().startswith("figura"):
                clean_caption = f"Figura 1 – {clean_caption}"
                
            run_fig_title = p_fig_title.add_run(clean_caption)
            run_fig_title.bold = True
            run_fig_title.font.size = Pt(10)
            run_fig_title.font.color.rgb = RGBColor(0, 0, 0)
            
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(2)
            p_img.paragraph_format.space_after = Pt(4)
            
            try:
                p_img.add_run().add_picture(image_path, width=Inches(5.5))
                
                # ABNT Source Attribution (Below Image)
                p_source = doc.add_paragraph()
                p_source.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_source.paragraph_format.space_before = Pt(2)
                p_source.paragraph_format.space_after = Pt(14)
                
                agency_name = str(config.get("agency_name") or "Secretaria Municipal de Trânsito").replace("|", "").strip()
                source_text = f"Fonte: {agency_name} / Motor Neuro-Simbólico CARINA (2026)."
                
                run_source = p_source.add_run(source_text)
                run_source.font.size = Pt(9)
                run_source.italic = True
                run_source.font.color.rgb = RGBColor(80, 80, 80)
            except Exception as e:
                logging.error(f"[STRUCTURED_REPORT_BUILDER] Failed to insert chart: {e}")
                err_text = get_translated(config, "structured_report.error_loading_chart", "[ERRO: Não foi possível carregar a imagem do gráfico.]")
                doc.add_paragraph(err_text)
        else:
            logging.info("[STRUCTURED_REPORT_BUILDER] Chart image path not found or empty. Adding placeholder.")
            heading = config.get("chart_title") or get_translated(config, "structured_report.xai_chart_title", "2. MAPA DE PLANEJAMENTO TÁTICO DA MALHA VIÁRIA")
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            run_h = h.add_run(heading.replace("|", "").strip())
            run_h.bold = True
            run_h.font.size = Pt(12)

            p_placeholder = doc.add_paragraph()
            p_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_placeholder.paragraph_format.space_before = Pt(10)
            p_placeholder.paragraph_format.space_after = Pt(10)
            r = p_placeholder.add_run("[Figura 1 – Mapa Espacial da Malha Viária: Renderização do Mapa Indisponível nesta Sessão]")
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(120, 120, 120)
