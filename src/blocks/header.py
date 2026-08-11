# CARINA (Controlled Artificial Road-traffic Intelligence Network Architecture) is an open-source AI ecosystem for real-time, adaptive control of urban traffic light networks.
# Copyright (C) 2026 Gabriel Moraes - Noxfort Systems
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# File: src/blocks/header.py
# Author: Gabriel Moraes
# Date: 2026-07-02

import os
import logging
from typing import Dict, Any

from .base import ReportBlock

try:
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

class HeaderBlock(ReportBlock):
    """Generates the organizational header including an optional logo and agency details."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        logo_path = config.get("logo_path")
        agency_name = str(config.get("agency_name", "Prefeitura Municipal / Secretaria de Trânsito"))
        department_name = str(config.get("department_name", "Departamento de Mobilidade Inteligente"))
        
        clean_agency = agency_name.replace("|", "").strip().upper()
        clean_dept = department_name.replace("|", "").strip()

        is_valid_logo = False
        if logo_path and os.path.exists(logo_path):
            ext = os.path.splitext(logo_path.lower())[1]
            if ext in ['.png', '.jpg', '.jpeg']:
                is_valid_logo = True
            try:
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                
                # Apply explicit widths: 3.0 cm for Logo cell (20%), 14.0 cm for Text cell (80%)
                table.columns[0].width = Cm(3.0)
                table.columns[1].width = Cm(14.0)

                # Left cell: Compact Logo
                cell_logo = table.cell(0, 0)
                cell_logo.width = Cm(3.0)
                p_logo = cell_logo.paragraphs[0]
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_logo.add_run().add_picture(logo_path, width=Cm(2.5))
                
                # Right cell: Wide, Centralized Agency & Department Text
                cell_text = table.cell(0, 1)
                cell_text.width = Cm(14.0)

                p_agency = cell_text.paragraphs[0]
                p_agency.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_agency.paragraph_format.space_before = Pt(0)
                p_agency.paragraph_format.space_after = Pt(2)
                
                run_agency = p_agency.add_run(clean_agency)
                run_agency.bold = True
                run_agency.font.size = Pt(11)
                run_agency.font.color.rgb = RGBColor(0, 0, 0)
                
                p_dept = cell_text.add_paragraph()
                p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_dept.paragraph_format.space_before = Pt(0)
                p_dept.paragraph_format.space_after = Pt(2)
                
                run_dept = p_dept.add_run(clean_dept)
                run_dept.font.size = Pt(9.5)
                run_dept.font.color.rgb = RGBColor(0, 0, 0)
                
            except Exception as e:
                logging.error(f"[STRUCTURED_REPORT_BUILDER] Failed to add header logo table: {e}")
        else:
            p_agency = doc.add_paragraph()
            p_agency.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_agency.paragraph_format.space_before = Pt(0)
            p_agency.paragraph_format.space_after = Pt(2)
            
            run_agency = p_agency.add_run(clean_agency)
            run_agency.bold = True
            run_agency.font.size = Pt(11)
            run_agency.font.color.rgb = RGBColor(0, 0, 0)
            
            p_dept = doc.add_paragraph()
            p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_dept.paragraph_format.space_before = Pt(0)
            p_dept.paragraph_format.space_after = Pt(2)
            
            run_dept = p_dept.add_run(clean_dept)
            run_dept.font.size = Pt(9.5)
            run_dept.font.color.rgb = RGBColor(0, 0, 0)
            
        p_div = doc.add_paragraph()
        p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_div.paragraph_format.space_before = Pt(4)
        p_div.paragraph_format.space_after = Pt(12)
