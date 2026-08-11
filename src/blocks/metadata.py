# CARINA (Controlled Artificial Road-traffic Intelligence Network Architecture) is an open-source AI ecosystem for real-time, adaptive control of urban traffic light networks.
# Copyright (C) 2026 Gabriel Moraes - Noxfort Systems
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as
# published by the Free Software Foundation, either version 3 of the
# License, or (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# File: src/blocks/metadata.py
# Author: Gabriel Moraes
# Date: 2026-07-02

from typing import Dict, Any

from .base import ReportBlock, get_translated

try:
    from docx.shared import Pt, RGBColor
except ImportError:
    pass

class MetadataBlock(ReportBlock):
    """Creates a structured metadata table with information regarding the agent or run under analysis."""
    def build(self, doc: Any, context: Dict[str, Any], config: Dict[str, Any]) -> None:
        agent_id = context.get("agent_id", "UNKNOWN")
        scenario = context.get("scenario", "Live Session")
        engine_ver = context.get("engine_version", "CARINA v1.0.0")
        
        mode = config.get("mode", "XAI")
        
        metadata_title = config.get("metadata_title")
        metadata_rows = context.get("metadata_rows") or config.get("metadata_rows")
        
        if metadata_title is not None and metadata_rows is not None:
            heading = metadata_title
            headers = metadata_rows
        elif mode == "MFD":
            heading = get_translated(config, "structured_report.mfd_metadata_title", "1. AMBIENTE OPERACIONAL E IDENTIFICAÇÃO")
            headers = [
                (get_translated(config, "structured_report.mfd_label_scenario", "Cenário de Simulação:"), scenario),
                (get_translated(config, "structured_report.mfd_label_engine", "Motor Analítico MFD:"), engine_ver),
                (get_translated(config, "structured_report.mfd_label_control", "Sistema de Controle:"), agent_id)
            ]
        else:
            heading = get_translated(config, "structured_report.xai_metadata_title", "1. AMBIENTE OPERACIONAL E IDENTIFICAÇÃO")
        clean_heading = str(heading).replace("|", "").strip()
        if not clean_heading.startswith("1."):
            clean_heading = f"1. {clean_heading}"
            
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        run_h = h.add_run(clean_heading)
        run_h.bold = True
        run_h.font.size = Pt(12)
        run_h.font.color.rgb = RGBColor(0, 0, 0)
        
        headers = [
            (get_translated(config, "structured_report.xai_label_agent", "Identificador do Agente:"), agent_id),
            (get_translated(config, "structured_report.xai_label_scenario", "Cenário de Operação:"), scenario),
            (get_translated(config, "structured_report.xai_label_engine", "Motor Analítico:"), engine_ver)
        ]
            
        table = doc.add_table(rows=len(headers), cols=2)
        table.style = 'Table Grid'
        
        for idx, (label, val) in enumerate(headers):
            clean_lbl = str(label).replace("|", "").strip()
            clean_val = str(val).replace("|", "").strip()
            
            cell_label = table.cell(idx, 0)
            run_lbl = cell_label.paragraphs[0].add_run(clean_lbl)
            run_lbl.bold = True
            run_lbl.font.size = Pt(10)
            run_lbl.font.color.rgb = RGBColor(0, 0, 0)
            
            cell_val = table.cell(idx, 1)
            run_val = cell_val.paragraphs[0].add_run(clean_val)
            run_val.font.size = Pt(10)
            run_val.font.color.rgb = RGBColor(0, 0, 0)
            
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
